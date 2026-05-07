# -*- coding: utf-8 -*-
"""
Market WatchDog 포트폴리오 문서 생성기
실행: python make_portfolio.py
출력: Market_WatchDog_Portfolio.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS = "screenshots"

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

GREEN = RGBColor(0x00, 0xC8, 0x5A)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x6B, 0x72, 0x80)


def heading(text, level=1, color=None):
    p = doc.add_paragraph(); p.clear()
    r = p.add_run(text); r.bold = True
    if level == 1:
        r.font.size = Pt(20); r.font.color.rgb = color or DARK
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), '00C85A')
        pBdr.append(bot); pPr.append(pBdr)
    elif level == 2:
        r.font.size = Pt(14); r.font.color.rgb = color or DARK
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        r.font.size = Pt(12); r.font.color.rgb = color or GRAY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p


def body(text, indent=False, color=None, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if bold:  r.bold = True
    if indent: p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    return p


def bullet(text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(10.5)
    if color: r.font.color.rgb = color
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    return p


def img_placeholder(label, width_cm=15.5, height_cm=8):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    tc   = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E5E7EB'); tcPr.append(shd)
    tr   = tbl.rows[0]._tr; trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'exact'); trPr.append(trH)
    p    = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r    = p.add_run(f"[스크린샷]  {label}")
    r.font.size = Pt(11); r.font.color.rgb = GRAY; r.bold = True
    doc.add_paragraph()
    return tbl


def add_image(filename, width_cm=15.5, caption=None):
    path = os.path.join(SCREENSHOTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(width_cm))
        if caption:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(caption); cr.font.size = Pt(9); cr.font.color.rgb = GRAY
        doc.add_paragraph()
    else:
        img_placeholder(caption or filename, width_cm=width_cm)


def kv_table(pairs):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(pairs):
        kr = tbl.cell(i, 0).paragraphs[0].add_run(k)
        vr = tbl.cell(i, 1).paragraphs[0].add_run(v)
        kr.bold = True; kr.font.size = Pt(10); kr.font.color.rgb = DARK
        vr.font.size = Pt(10)
        tc   = tbl.cell(i, 0)._tc; tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F3F4F6'); tcPr.append(shd)
    doc.add_paragraph()


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(10); r.font.color.rgb = GREEN
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════
#  표지
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("Market WatchDog")
r.bold = True; r.font.size = Pt(34); r.font.color.rgb = DARK

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("한국 단기금융시장 모니터링 대시보드")
r.font.size = Pt(14); r.font.color.rgb = GRAY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
r = p.add_run("CP  ·  ABCP  ·  AB단기사채  |  ML 이상탐지  |  Flask + SQLite + OpenAI")
r.font.size = Pt(11); r.font.color.rgb = GREEN; r.bold = True

add_image("cover.png", caption="Market WatchDog 대시보드")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("2026.05  |  Portfolio Document")
r.font.size = Pt(10); r.font.color.rgb = GRAY

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  1. 프로젝트 개요
# ════════════════════════════════════════════════════════════
heading("1. 프로젝트 개요", 1)
body(
    "Market WatchDog은 국내 단기금융시장(CP·ABCP·AB단기사채)의 발행·잔액·신용등급·만기구조 "
    "데이터를 월별로 수집·가공하고, 머신러닝 기반 이상탐지 모델로 시장 스트레스를 정량화해 "
    "실시간으로 시각화하는 Flask 기반 웹 대시보드입니다.", size=11
)
body(
    "레고랜드 사태(2022.10), 새마을금고 사태(2023.07) 등 주요 신용 이벤트가 CP·ABCP 시장에서 "
    "먼저 신호를 보낸다는 점에 착안해, 정량적·자동화된 스트레스 탐지 도구를 직접 설계·구현했습니다.", size=11
)

heading("기본 정보", 2)
kv_table([
    ("프로젝트명",  "Market WatchDog"),
    ("개발 기간",   "2025년 하반기 ~ 2026년 5월 (지속 업데이트)"),
    ("개발 목적",   "단기금융시장 스트레스를 정량화해 담당자의 의사결정을 보조하는 대시보드 구현"),
    ("데이터 범위", "2016년 ~ 2026년 4월 (월별, CP·ABCP·AB단기사채)"),
    ("기술 스택",   "Python 3.12 / Flask / SQLite / Pandas / scikit-learn / Chart.js / OpenAI"),
    ("데이터 원천", "한국예탁결제원(Seibro) 발행·잔액 데이터"),
    ("보안 레이어", "Date-Based Dynamic Decoy Security Layer (에니그마 컨셉 기반 자체 설계)"),
])


# ════════════════════════════════════════════════════════════
#  2. 시스템 아키텍처
# ════════════════════════════════════════════════════════════
heading("2. 시스템 아키텍처", 1)

heading("데이터 파이프라인 (월 1회 자동 실행)", 2)
body("① 원천 데이터 수집", bold=True)
bullet("발행 데이터: Selenium 기반 Seibro 크롤러 → CP/ABCP/AB단기사채 월별 CSV")
bullet("잔액 데이터: Seibro 잔액 크롤러 → 종목별 월말 잔액 CSV")

body("② 전처리 자동화", bold=True)
bullet("신용등급 정규화, 만기 버킷 분류(ultra~xlong), 월별 누적 시계열 생성")

body("③ DB 적재", bold=True)
bullet("전처리 CSV → SQLite(market_watchdog.db) upsert")

body("④ ML 모델 적용", bold=True)
bullet("Model 1: 12개월 슬라이딩 윈도우 → 롤오버 비율 추세 이탈 Z-score 산출")
bullet("Model 2: Isolation Forest → 다변량 이상 스코어(0~1) 산출")
bullet("출력: model_output.json")

body("⑤ AI 요약 생성", bold=True)
bullet("model_output.json 기반 → GPT-4o-mini → ai_summary.json")

body("전체 실행 명령", bold=True)
code("python run_pipeline.py 26년4월")

heading("구조 선택의 이유", 2)
body(
    "각 설계 결정은 '현재 단계에서 가장 적절한 선택'을 기준으로 했습니다. "
    "아래는 주요 의사결정과 그 근거입니다.", size=10.5
)
kv_table([
    ("SQLite 선택",
     "월별 갱신·단일 사용자·로컬 실행 환경에서 서버형 DB는 과합니다. "
     "파일 하나로 배포 가능하며, 필요 시 PostgreSQL로 전환이 용이하도록 ORM 없이 sqlite3를 직접 사용했습니다."),
    ("모델 결과를 JSON 캐싱",
     "데이터가 월 1회 바뀌므로 요청마다 ML 모델을 실행할 이유가 없습니다. "
     "파이프라인에서 한 번만 실행 후 model_output.json에 저장하고, Flask는 JSON을 읽어 서빙합니다. "
     "응답 속도와 API 비용을 동시에 절약합니다."),
    ("AI 요약을 파이프라인에 포함",
     "페이지 로드마다 AI API를 호출하면 비용이 선형으로 증가합니다. "
     "데이터 갱신 주기와 요약 생성 주기를 일치시켜 월 1회만 호출합니다. "
     "긴급 재생성은 관리자 페이지 버튼으로 온디맨드 처리합니다."),
    ("두 ML 모델 결합",
     "Isolation Forest 단독은 다변량 이상 탐지에는 강하지만 방향성(유입/유출)을 알 수 없습니다. "
     "Sliding Window Z-score는 추세 방향을 보여주지만 다변량 패턴을 잡지 못합니다. "
     "두 모델을 결합해 서로의 약점을 보완했습니다."),
    ("Decoy 레이어 설계 방향",
     "표준 암호화(AES-GCM)를 대체하는 것이 아닙니다. "
     "비인증 접근 시 접근 사실을 모르게 하면서 로그를 남기는 탐지·지연·추적 목적의 별도 레이어입니다. "
     "강력한 보안보다 '보안 레이어를 설계·구현할 수 있다'는 것을 보이는 데 초점을 뒀습니다."),
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  3. ML 모델 설계
# ════════════════════════════════════════════════════════════
heading("3. ML 모델 설계", 1)

heading("Model 1 — 슬라이딩 윈도우 추세 이탈 탐지", 2)
body("직전 12개월 롤오버 비율(발행액/잔액)의 평균·표준편차를 기준선으로 삼아 이탈 크기를 Z-score로 수치화합니다.")
kv_table([
    ("입력 변수", "롤오버 비율, 단기만기 비중, 고신용 발행 비중"),
    ("출력",      "sw_z (이탈 크기), sw_signal (정상/주의/경고)"),
    ("판정 기준", "|z| ≥ 2.0 → 경고,  |z| ≥ 1.5 → 주의"),
])

heading("Model 2 — Isolation Forest 다변량 이상 탐지", 2)
body("롤오버 비율·가속도·신용등급 변화율·만기구조 변화 4개 피처를 Isolation Forest로 학습합니다.")
kv_table([
    ("입력 변수",   "rollover_ratio, rollover_accel, grade_change, tenor_change"),
    ("알고리즘",    "Isolation Forest (scikit-learn)"),
    ("학습 데이터", "2016.01 ~ 2024.12"),
    ("출력",        "anomaly_score (0=정상, 1=최고위험)"),
    ("저장",        "model/*.pkl (joblib) — 매월 apply_saved_model()로 재사용"),
])

heading("레벨 판정 (Combined)", 2)
kv_table([
    ("정상", "anomaly_score < 0.4  AND  sw_signal = 정상"),
    ("주의", "anomaly_score ≥ 0.4  OR   sw_signal = 주의"),
    ("경고", "anomaly_score ≥ 0.6  OR   sw_signal = 경고"),
    ("위험", "anomaly_score ≥ 0.8"),
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  4. 보안 레이어
# ════════════════════════════════════════════════════════════
heading("4. Date-Based Dynamic Decoy Security Layer", 1)
body(
    "비인증 API 접근 또는 파일 탈취 시 원본 데이터 대신 '미끼 데이터(Decoy)'를 반환하는 "
    "자체 설계 보안 레이어입니다. 실제 암호화의 대체재가 아닌 접근 감지·유출 지연·로그 추적을 위한 방어층입니다.", size=11
)

heading("설계 컨셉 — 에니그마 + 도둑맞은 편지", 2)
bullet("에니그마: 월별 업무 메일을 Seed 문서(플러그보드 배선)로 사용, 날짜 자릿수 합산을 로터 회전으로 활용")
bullet("도둑맞은 편지: Seed 문서는 mail_archive/ 12개 업무 메일 중 하나 — 공개된 장소에 숨김")
bullet("공격자가 알고리즘(코드)을 알아도 어느 메일이 Seed인지 모르면 해독 불가 (26! 경우의 수)")

heading("치환 알고리즘", 2)
body("① Seed 문서에서 알파벳 첫 등장 순서 26자 추출 → 치환 알파벳 생성", indent=True)
body("② 접근 날짜의 월/일 자릿수 합산 → 회전 칸수 결정", indent=True)
code("예: 5월 6일 → 5+6 = 11칸 회전")
body("③ JSON 키 + 값 전부 치환 (숫자→알파벳, 알파벳→알파벳, 한글→[X])", indent=True)

heading("치환 결과 예시", 2)
kv_table([
    ("원본",  '{"anomaly_score": 0.312, "final_level": "정상", "YYYYMM": 202604}'),
    ("Decoy", '{"txwqtvn_rcwfg": "T.ELC", "hyxtv_vgjgv": "[X]", "nnnnqq": "CTCPTG"}'),
])

heading("관리자 콘솔 (/admin)", 2)
bullet("Flask 세션 기반 로그인")
bullet("mail_archive/ 12개 업무 메일 카드 → 클릭으로 Seed 즉시 교체")
bullet("접근 로그 실시간 조회 (NORMAL_ACCESS / DECOY_RETURNED / UNAUTHORIZED)")

add_image("admin.png", caption="관리자 콘솔 — Seed 선택 + 접근 통계 + 로그")
add_image("decoy_demo.png", caption="Decoy 데모 — 좌: 원본 데이터, 우: 치환된 Decoy 데이터")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  5. 주요 기능
# ════════════════════════════════════════════════════════════
heading("5. 주요 기능", 1)

heading("5-1. INDICATORS — 시계열 지표 대시보드", 2)
body("CP·ABCP·AB단기사채 3개 시장의 잔액·발행액·만기구조·신용등급 비중을 2016년부터 월별 차트로 제공합니다.")
bullet("기간 선택: 전체 / 최근 24·12·6·3개월 / 커스텀 범위(YYYYMM 직접 입력)")
bullet("잔액 추이: 10년 시계열, 주요 사건 세로선 자동 표시")
bullet("만기구조: ultra(≤7일) ~ xlong(366일+) 6버킷 스택 차트")
bullet("신용등급 비중: top(A1) / high(A2) / mid(A3) / low(B이하) 4버킷")
add_image("indicators.png", caption="INDICATORS — CP 잔액·발행액·만기구조·신용등급 차트")

doc.add_page_break()

heading("5-2. 시장감시 — 시장 이상 감지 대시보드", 2)
body(
    "슬라이딩 윈도우 + Isolation Forest 결합 결과를 실시간 시각화합니다. "
    "시스템이 결론을 내리는 것이 아니라, 담당자가 수치를 보고 직접 판단할 수 있도록 "
    "보조 지표를 제공하는 구조입니다."
)
bullet("레벨 카드: CP·ABCP·AB단기사채 각각 정상/주의/경고/위험 색상 표시")
bullet("시장 스트레스 시계열: 경보(0.6)·주의(0.4) 기준선 포함")
bullet("추세 이탈 시계열: 경보(±2.0)·주의(±1.5) 기준선, 탭 전환으로 전환")
bullet("날짜 범위 선택 + 6·12·24개월·전체 테이블 필터")
bullet("AI 시장 해석 카드: GPT-4o-mini가 현재 시장을 비전문가 언어로 자동 요약")
add_image("analysis_stress.png", caption="시장감시 — 레벨 카드 + 시장 스트레스 시계열 + AI 해석")
add_image("analysis_trend.png", caption="시장감시 — 추세 이탈 시계열 (탭 전환)")

doc.add_page_break()

heading("5-3. AI 시장 도우미 — 실시간 채팅", 2)
body(
    "시장감시 페이지 우측 하단 채팅 버튼을 누르면 AI 도우미와 대화할 수 있습니다. "
    "현재 CP·ABCP·AB단기사채 데이터가 자동으로 컨텍스트에 주입되어 시장 데이터 기반 구체적인 답변을 제공합니다."
)
bullet("모델: OpenAI GPT-4o-mini")
bullet("컨텍스트: 최신 model_output.json 기준 레벨·스트레스·추세이탈 자동 포함")
bullet("대화 이력 유지: 최근 10턴 유지, 연속적인 질의응답 가능")
bullet("비전문가 대상: 금융 용어를 쉬운 말로 풀어 설명하도록 시스템 프롬프트 설계")

body("실제 대화 예시", bold=True)
kv_table([
    ("사용자", "현재 시장동향 어때?"),
    ("AI",     "현재 단기채권 시장은 전반적으로 정상적인 상황이에요. 하지만 ABCP(자산담보기업어음) 중 "
               "일부에서 '경고' 신호가 나타나고 있고, 몇몇 채권은 '주의' 상태에 있어요."),
    ("사용자", "ABCP는 연속적으로 위험하다고 나오는데 얼마나 지속되면 문제가 생기나?"),
    ("AI",     "ABCP가 계속 위험 신호를 보이면 조만간 자금 조달에 어려움이 생길 수 있어요. "
               "몇 주에서 몇 달 안에 확인해야 하는데, 시장 상황에 따라 다를 수 있으니 계속 지켜보는 게 중요해요."),
    ("사용자", "과거 사건에 대입해 보면 현재 상황이 비슷한 케이스가 있나?"),
    ("AI",     "2020년 중반 팬데믹 시기처럼 특정 채권에서 스트레스와 주의 신호가 동시에 나타나는 패턴과 "
               "유사해요. 그때도 시장이 신중하게 반응했던 만큼 계속 예의주시하는 것이 중요해요."),
])
add_image("ai_chat.png", caption="AI 시장 도우미 — 시장 데이터 기반 실시간 Q&A")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  6. 기술적 특징
# ════════════════════════════════════════════════════════════
heading("6. 기술적 특징", 1)

heading("데이터 엔지니어링", 2)
bullet("Selenium + WebDriverManager 기반 Seibro 자동 크롤링")
bullet("다양한 인코딩(UTF-8-sig / CP949) 자동 감지 CSV 로더")
bullet("신용등급 정규화: 전각문자·특수기호 포함 원천 → A1/A2+/A2/… 표준 코드")
bullet("월별 incremental append: 이전 누적본 + 신규 1개월 → 중복 제거 후 저장")

heading("ML 파이프라인", 2)
bullet("학습은 최초 1회, 이후 apply_saved_model()로 pkl 재사용 → 재학습 없이 월별 갱신")
bullet("Z-score 기반 추세 이탈 + 다변량 Isolation Forest 결합으로 오탐 감소")

heading("Flask API 설계", 2)
bullet("20개 이상 REST 엔드포인트: /api/{cp|abcp|ab}/{balance|issue|tenor|grade}")
bullet("?from=YYYYMM&to=YYYYMM 파라미터 기간 필터 지원")
bullet("/api/secure/model/latest: Decoy 보안 레이어 적용 엔드포인트")
bullet("Flask 세션 기반 관리자 인증 (/login, /logout, /admin)")

heading("프론트엔드", 2)
bullet("순수 Chart.js 4.x — 별도 프레임워크 없이 다수 차트 렌더링")
bullet("탭 전환(시장 스트레스 / 추세 이탈), 날짜 범위 선택, 테이블 기간 필터")
bullet("범례에 주의·경보 기준선 설명, 탭 전환 시 라벨 자동 변경")
bullet("AI 채팅 플로팅 버튼, 채팅 이력 유지, 시장 데이터 컨텍스트 자동 주입")


# ════════════════════════════════════════════════════════════
#  7. 파일 구조
# ════════════════════════════════════════════════════════════
heading("7. 주요 파일 구조", 1)
kv_table([
    ("app.py",                           "Flask 서버 · REST API · 보안 레이어 · 관리자 라우트"),
    ("run_pipeline.py",                  "전체 파이프라인 원커맨드 실행기"),
    ("monthly_preprocess_automation.py", "월별 원천 → 표준 CSV 전처리 자동화"),
    ("db_upsert_patch.py",               "전처리 CSV → market_watchdog.db upsert"),
    ("crawler.py",                       "연합인포맥스 뉴스 크롤러"),
    ("model/model1.py",                  "슬라이딩 윈도우 추세 이탈 탐지"),
    ("model/model2.py",                  "Isolation Forest 이상 탐지 (학습)"),
    ("model/combined.py",                "두 모델 결합 + JSON 출력 + pkl 저장"),
    ("security/cipher.py",               "Seed 기반 치환 암호 엔진"),
    ("security/seed_manager.py",         "Seed 문서 DB 관리 · 접근 로그"),
    ("security/auth.py",                 "JWT 기반 API 인증"),
    ("mail_archive/",                    "Seed 후보 업무 메일 12개"),
    ("templates/analysis.html",          "시장감시 대시보드 (차트 + AI 채팅)"),
    ("templates/admin.html",             "관리자 콘솔 (Seed 관리 + 접근 로그)"),
    ("templates/demo.html",              "Decoy 보안 레이어 시연 페이지"),
    (".env.example",                     "환경변수 설정 템플릿 (API 키는 직접 입력)"),
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  8. 성과 및 의의
# ════════════════════════════════════════════════════════════
heading("8. 프로젝트 성과 및 의의", 1)
bullet("2016~2026년 약 120개월 CP·ABCP·AB단기사채 통합 시계열 구축", color=DARK)
bullet("크롤링 → 전처리 → ML → 시각화 전 과정을 단일 파이프라인으로 자동화", color=DARK)
bullet("에니그마 컨셉 기반 자체 설계 보안 레이어 — Seed 없이 해독 불가 (26! 경우의 수)", color=DARK)
bullet("OpenAI GPT-4o-mini 연동으로 비전문가 대상 자동 시장 해석 및 실시간 Q&A", color=DARK)
bullet("Flask 세션 인증 + 관리자 콘솔로 운영 가능한 완성도 구현", color=DARK)

heading("한계 및 향후 개선 방향", 2)
bullet("Seibro 직접 API 연동 시 크롤링 없이 완전 자동화 가능")
bullet("Seed 문서 DB 저장 → 별도 파일 분리 보관으로 보안 강화 여지")
bullet("클라우드 배포 (현재 로컬 서버) → AWS/GCP 배포 예정")

heading("부록. 스크린샷 저장 가이드", 1)
body("screenshots/ 폴더에 아래 파일명으로 저장 후 python make_portfolio.py 재실행하면 자동 삽입됩니다.", size=10, color=GRAY)
kv_table([
    ("cover.png",           "대시보드 전체 화면 (표지용)"),
    ("indicators.png",      "INDICATORS — CP 차트 전체"),
    ("analysis_stress.png", "시장감시 — 레벨 카드 + 시장 스트레스 차트"),
    ("analysis_trend.png",  "시장감시 — 추세 이탈 탭"),
    ("ai_chat.png",         "AI 채팅 화면"),
    ("decoy_demo.png",      "Decoy 데모 좌우 비교"),
    ("admin.png",           "관리자 콘솔"),
])


# ── 저장 ─────────────────────────────────────────────────────
out = "Market_WatchDog_Portfolio.docx"
doc.save(out)
print(f"[완료] {out} 저장됨")
